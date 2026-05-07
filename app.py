import os
import json
import csv
from datetime import datetime, timezone, timedelta
from functools import wraps
from io import StringIO
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_from_directory, Response
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import io
from fpdf import FPDF
from docx import Document
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-replace-in-production')

# DB Configuration — always use DATABASE_URL if set (Railway PostgreSQL), else local SQLite
db_url = os.environ.get('DATABASE_URL', 'sqlite:///reports.db')
if db_url.startswith('postgres://'):
    db_url = db_url.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = db_url

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

UPLOAD_FOLDER = os.path.join(app.root_path, 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024 * 10 # 50 MB total to allow multiple files
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'pdf', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'txt'}

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message_category = 'info'

# ROLES AND PERMISSIONS
ROLES = {
    'admin': ['view_all_reports', 'view_dept_reports', 'view_own_reports', 'submit_report', 'view_analytics', 'manage_users', 'manage_departments', 'export_data', 'delete_reports', 'approve_reports', 'manage_performance', 'manage_holidays', 'manage_kpis', 'manage_leaves'],
    'manager': ['view_dept_reports', 'view_analytics', 'export_data', 'approve_reports', 'manage_performance', 'manage_kpis', 'manage_leaves'],
    'employee': ['view_own_reports', 'submit_report', 'manage_leaves'],
    'viewer': ['view_all_reports', 'view_analytics']
}

# --- MODELS ---
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(256), nullable=False)
    full_name = db.Column(db.String(150))
    role = db.Column(db.String(50), default='employee')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    def has_permission(self, permission):
        # Check built-in roles first
        if self.role in ROLES:
            return permission in ROLES.get(self.role, [])
        # Check custom roles from database
        custom_role = CustomRole.query.filter_by(name=self.role).first()
        if custom_role:
            perms = json.loads(custom_role.permissions or '[]')
            return permission in perms
        return False

class Department(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)
    description = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

class Report(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    employee_name = db.Column(db.String(150), nullable=False)
    department = db.Column(db.String(100), nullable=False)
    report_text = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(50), default='Submitted')
    images = db.Column(db.Text) # JSON string array
    date_submitted = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    user = db.relationship('User', backref=db.backref('reports', lazy=True))

class Holiday(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    date = db.Column(db.Date, unique=True, nullable=False)
    description = db.Column(db.String(200))

class PerformanceNote(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    admin_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    entry_type = db.Column(db.String(50), default='Improvement') # Improvement, Win
    description = db.Column(db.Text, nullable=False)
    date_logged = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    
    employee = db.relationship('User', foreign_keys=[employee_id], backref=db.backref('performance_notes', lazy=True))
    admin = db.relationship('User', foreign_keys=[admin_id])

class KPIDefinition(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    department_id = db.Column(db.Integer, db.ForeignKey('department.id'), nullable=False)
    metric_name = db.Column(db.String(150), nullable=False)
    metric_type = db.Column(db.String(50), default='number') # number, percentage, boolean
    target_value = db.Column(db.Float, nullable=True)
    description = db.Column(db.Text, nullable=True)
    
    department = db.relationship('Department', backref=db.backref('kpis', lazy=True))

class DailyKPILog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    manager_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    date = db.Column(db.Date, nullable=False, default=lambda: datetime.now(timezone.utc).date())
    kpi_definition_id = db.Column(db.Integer, db.ForeignKey('kpi_definition.id'), nullable=False)
    actual_value = db.Column(db.Float, nullable=False)
    
    employee = db.relationship('User', foreign_keys=[employee_id], backref=db.backref('kpi_logs', lazy=True))
    manager = db.relationship('User', foreign_keys=[manager_id])
    kpi_definition = db.relationship('KPIDefinition')

class LeaveRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    start_date = db.Column(db.Date, nullable=False)
    end_date = db.Column(db.Date, nullable=False)
    reason = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(50), default='Pending') # Pending, Approved, Rejected
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    
    user = db.relationship('User', backref=db.backref('leave_requests', lazy=True))

class CustomRole(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)
    description = db.Column(db.String(200))
    permissions = db.Column(db.Text, default='[]')  # JSON list of permission strings
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

def permission_required(permission):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login'))
            if not current_user.has_permission(permission):
                flash("You don't have permission to access this page.", "danger")
                return redirect(url_for('index')), 403
            return f(*args, **kwargs)
        return decorated_function
    return decorator

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- ROUTES ---

@app.template_filter('from_json')
def from_json_filter(value):
    try:
        return json.loads(value or '[]')
    except Exception:
        return []

@app.route('/health')
def health():
    return jsonify({'status': 'healthy'}), 200

@app.context_processor
def inject_now():
    return {'now': datetime.now(timezone.utc)}

@app.route('/')
@login_required
def index():
    now_utc = datetime.now(timezone.utc)
    today_dt = now_utc.date()
    yesterday_dt = today_dt - timedelta(days=1)

    # Helper for date ranges (Postgres safe)
    def get_day_range(d):
        start = datetime.combine(d, datetime.min.time()).replace(tzinfo=timezone.utc)
        end = datetime.combine(d, datetime.max.time()).replace(tzinfo=timezone.utc)
        return start, end

    t_start, t_end = get_day_range(today_dt)
    y_start, y_end = get_day_range(yesterday_dt)

    if current_user.has_permission('view_all_reports') or current_user.has_permission('view_dept_reports'):
        total_reports = Report.query.count()
        yesterday_total = Report.query.filter(Report.date_submitted >= y_start, Report.date_submitted <= y_end).count()
        today_reports = Report.query.filter(Report.date_submitted >= t_start, Report.date_submitted <= t_end).count()
        yesterday_today = yesterday_total # Legacy variable name fix
        unique_employees = db.session.query(db.func.count(db.distinct(Report.user_id))).scalar() or 0
        total_departments = Department.query.count()
        yesterday_employees = unique_employees # Fallback
    else:
        total_reports = Report.query.filter_by(user_id=current_user.id).count()
        yesterday_total = Report.query.filter_by(user_id=current_user.id).filter(Report.date_submitted >= y_start, Report.date_submitted <= y_end).count()
        today_reports = Report.query.filter_by(user_id=current_user.id).filter(Report.date_submitted >= t_start, Report.date_submitted <= t_end).count()
        yesterday_today = yesterday_total
        unique_employees = 1
        yesterday_employees = 1
        total_departments = 0

    is_holiday = Holiday.query.filter_by(date=today_dt).first() is not None
    due_employees = []
    is_due = False

    if not is_holiday:
        submitted_today = set(r.user_id for r in Report.query.filter(Report.date_submitted >= t_start, Report.date_submitted <= t_end).all())
        
        # Check leaves
        active_leaves = LeaveRequest.query.filter(
            LeaveRequest.status == 'Approved',
            LeaveRequest.start_date <= today_dt,
            LeaveRequest.end_date >= today_dt
        ).all()
        on_leave_users = set(l.user_id for l in active_leaves)
        
        if current_user.has_permission('view_all_reports') or current_user.has_permission('view_dept_reports'):
            employees = User.query.filter_by(role='employee').all()
            for emp in employees:
                if emp.id not in submitted_today and emp.id not in on_leave_users:
                    due_employees.append(emp)
        elif current_user.role == 'employee':
            if current_user.id not in submitted_today and current_user.id not in on_leave_users:
                is_due = True

    if current_user.has_permission('view_all_reports') or current_user.has_permission('view_dept_reports'):
        recent_activity = Report.query.order_by(Report.date_submitted.desc()).limit(5).all()
    else:
        recent_activity = Report.query.filter_by(user_id=current_user.id).order_by(Report.date_submitted.desc()).limit(5).all()

    # 7-day chart data
    chart_labels = []
    chart_data = []
    for i in range(6, -1, -1):
        day = today_dt - timedelta(days=i)
        ds, de = get_day_range(day)
        if current_user.has_permission('view_all_reports') or current_user.has_permission('view_dept_reports'):
            count = Report.query.filter(Report.date_submitted >= ds, Report.date_submitted <= de).count()
        else:
            count = Report.query.filter_by(user_id=current_user.id).filter(Report.date_submitted >= ds, Report.date_submitted <= de).count()
        chart_labels.append(day.strftime('%b %d'))
        chart_data.append(count)

    # Department performance
    dept_performance = []
    if current_user.has_permission('view_all_reports') or current_user.has_permission('view_dept_reports'):
        first_day = today_dt.replace(day=1)
        fd_start = datetime.combine(first_day, datetime.min.time()).replace(tzinfo=timezone.utc)
        
        last_month_end = first_day - timedelta(days=1)
        last_month_start = last_month_end.replace(day=1)
        lms_start = datetime.combine(last_month_start, datetime.min.time()).replace(tzinfo=timezone.utc)
        lme_end = datetime.combine(last_month_end, datetime.max.time()).replace(tzinfo=timezone.utc)

        depts = Department.query.all()
        for dept in depts:
            this_month = Report.query.filter(Report.department == dept.name, Report.date_submitted >= fd_start).count()
            last_month = Report.query.filter(Report.department == dept.name, Report.date_submitted >= lms_start, Report.date_submitted <= lme_end).count()
            
            if last_month > 0:
                change = round(((this_month - last_month) / last_month) * 100)
            else:
                change = 100 if this_month > 0 else 0
            dept_performance.append({'name': dept.name, 'count': this_month, 'change': change})
        dept_performance.sort(key=lambda x: x['count'], reverse=True)

    return render_template('index.html',
                           total_reports=total_reports, yesterday_total=yesterday_total,
                           today_reports=today_reports, yesterday_today=yesterday_today,
                           unique_employees=unique_employees, yesterday_employees=yesterday_employees,
                           total_departments=total_departments,
                           is_holiday=is_holiday, due_employees=due_employees, is_due=is_due,
                           recent_activity=recent_activity,
                           chart_labels=chart_labels, chart_data=chart_data,
                           dept_performance=dept_performance)

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        full_name = request.form.get('full_name')

        if User.query.filter_by(email=email).first() or User.query.filter_by(username=username).first():
            flash('Email or username already exists.', 'danger')
            return redirect(url_for('signup'))
        
        is_first_user = User.query.count() == 0
        role = 'admin' if is_first_user else 'employee'

        new_user = User(
            username=username, 
            email=email, 
            password=generate_password_hash(password),
            full_name=full_name,
            role=role
        )
        db.session.add(new_user)
        db.session.commit()
        
        login_user(new_user)
        flash('Account created successfully!', 'success')
        return redirect(url_for('index'))
    return render_template('signup.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        user = User.query.filter_by(email=email).first()
        
        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('index'))
        else:
            flash('Login failed. Check email and password.', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/submit', methods=['GET', 'POST'])
@permission_required('submit_report')
def submit():
    departments = Department.query.all()
    if request.method == 'POST':
        employee_name = request.form.get('employee_name')
        department = request.form.get('department')
        report_text = request.form.get('report_text')
        
        uploaded_images = request.files.getlist('images')
        saved_images = []
        
        for file in uploaded_images:
            if file and file.filename != '' and allowed_file(file.filename):
                file.seek(0, os.SEEK_END)
                size = file.tell()
                file.seek(0)
                if size > 5 * 1024 * 1024:
                    flash(f'File {file.filename} exceeds 5MB limit.', 'danger')
                    continue
                filename = secure_filename(f"{int(datetime.now().timestamp())}_{file.filename}")
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                saved_images.append(filename)

        report = Report(
            user_id=current_user.id,
            employee_name=employee_name,
            department=department,
            report_text=report_text,
            images=json.dumps(saved_images)
        )
        db.session.add(report)
        db.session.commit()
        flash('Report submitted successfully!', 'success')
        return redirect(url_for('view_reports'))
        
    return render_template('submit.html', departments=departments)

@app.route('/view')
@login_required
def view_reports():
    page = request.args.get('page', 1, type=int)
    department = request.args.get('department', '')
    employee_name = request.args.get('employee', '')
    date_filter = request.args.get('date', '')

    query = Report.query
    
    if not current_user.has_permission('view_all_reports'):
        if current_user.has_permission('view_dept_reports'):
             pass
        else:
             query = query.filter_by(user_id=current_user.id)

    if department:
        query = query.filter(Report.department.ilike(f'%{department}%'))
    if employee_name:
        query = query.filter(Report.employee_name.ilike(f'%{employee_name}%'))
    if date_filter:
        try:
            filter_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
            query = query.filter(db.func.date(Report.date_submitted) == filter_date)
        except ValueError:
            pass

    pagination = query.order_by(Report.date_submitted.desc()).paginate(page=page, per_page=10, error_out=False)
    departments = Department.query.all()
    
    return render_template('view.html', pagination=pagination, departments=departments, 
                           request_args=request.args, json=json)

@app.route('/report/<int:id>')
@login_required
def single_report(id):
    report = db.session.get(Report, id)
    if not report:
        flash('Report not found.', 'danger')
        return redirect(url_for('view_reports'))
        
    if not current_user.has_permission('view_all_reports') and not current_user.has_permission('view_dept_reports'):
        if report.user_id != current_user.id:
            flash('Permission denied.', 'danger')
            return redirect(url_for('view_reports'))
            
    images = json.loads(report.images) if report.images else []
    return render_template('single_report.html', report=report, images=images)

@app.route('/report/<int:id>/delete', methods=['POST'])
@permission_required('delete_reports')
def delete_report(id):
    report = db.session.get(Report, id)
    if report:
        if report.images:
            images = json.loads(report.images)
            for img in images:
                try:
                    os.remove(os.path.join(app.config['UPLOAD_FOLDER'], img))
                except OSError:
                    pass
        db.session.delete(report)
        db.session.commit()
        flash('Report deleted.', 'success')
    return redirect(url_for('view_reports'))

@app.route('/report/<int:id>/status', methods=['POST'])
@permission_required('approve_reports')
def update_report_status(id):
    report = db.session.get(Report, id)
    if not report:
        flash('Report not found.', 'danger')
        return redirect(url_for('view_reports'))
    
    new_status = request.form.get('status')
    if new_status in ['Approved', 'Rejected', 'Submitted']:
        report.status = new_status
        db.session.commit()
        flash(f'Report #{report.id} marked as {new_status}.', 'success')
    else:
        flash('Invalid status.', 'danger')
        
    return redirect(request.referrer or url_for('view_reports'))

@app.route('/uploads/<filename>')
@login_required
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/analytics')
@permission_required('view_analytics')
def analytics():
    dept_counts = db.session.query(Report.department, db.func.count(Report.id)).group_by(Report.department).all()
    dept_labels = [d[0] for d in dept_counts]
    dept_data = [d[1] for d in dept_counts]

    emp_counts = db.session.query(Report.employee_name, db.func.count(Report.id)).group_by(Report.employee_name).order_by(db.func.count(Report.id).desc()).limit(10).all()
    emp_labels = [e[0] for e in emp_counts]
    emp_data = [e[1] for e in emp_counts]

    seven_days_ago = datetime.now(timezone.utc) - timedelta(days=7)
    recent_reports = Report.query.filter(Report.date_submitted >= seven_days_ago).all()
    
    trend_dict = {}
    for i in range(7, -1, -1):
        d = (datetime.now(timezone.utc) - timedelta(days=i)).date().strftime('%Y-%m-%d')
        trend_dict[d] = 0
        
    for r in recent_reports:
        d_str = r.date_submitted.date().strftime('%Y-%m-%d')
        if d_str in trend_dict:
            trend_dict[d_str] += 1
            
    trend_labels = list(trend_dict.keys())
    trend_data = [trend_dict[k] for k in trend_labels]

    # Fetch all employees for the dropdown
    employees = User.query.filter_by(role='employee').order_by(User.full_name).all()
    employee_names = [e.full_name or e.username for e in employees]

    # Company-wide KPI Summary (Average of actual vs target for last 30 days)
    thirty_days_ago = datetime.now(timezone.utc).date() - timedelta(days=30)
    company_kpis = db.session.query(
        KPIDefinition.metric_name,
        db.func.avg(DailyKPILog.actual_value),
        KPIDefinition.target_value
    ).join(DailyKPILog, DailyKPILog.kpi_definition_id == KPIDefinition.id)\
     .filter(DailyKPILog.date >= thirty_days_ago)\
     .group_by(KPIDefinition.metric_name, KPIDefinition.target_value).all()
    
    kpi_summary = []
    for name, avg_val, target in company_kpis:
        kpi_summary.append({
            'name': name,
            'average': round(avg_val, 2),
            'target': target
        })

    return render_template('analytics.html', 
                           dept_labels=json.dumps(dept_labels), dept_data=json.dumps(dept_data),
                           emp_labels=json.dumps(emp_labels), emp_data=json.dumps(emp_data),
                           trend_labels=json.dumps(trend_labels), trend_data=json.dumps(trend_data),
                           employee_names=employee_names,
                           kpi_summary=kpi_summary)

@app.route('/export')
@permission_required('export_data')
def export():
    department = request.args.get('department', '')
    query = Report.query
    if department:
        query = query.filter(Report.department.ilike(f'%{department}%'))
    
    reports = query.order_by(Report.date_submitted.desc()).all()
    
    def generate():
        data = StringIO()
        writer = csv.writer(data)
        writer.writerow(['ID', 'Employee', 'Department', 'Report Text', 'Status', 'Images', 'Date Submitted'])
        yield data.getvalue()
        data.seek(0)
        data.truncate(0)
        for r in reports:
            images = ", ".join(json.loads(r.images)) if r.images else ""
            writer.writerow([r.id, r.employee_name, r.department, r.report_text, r.status, images, r.date_submitted.strftime('%Y-%m-%d %H:%M:%S')])
            yield data.getvalue()
            data.seek(0)
            data.truncate(0)
            
    response = Response(generate(), mimetype='text/csv')
    response.headers.set('Content-Disposition', f'attachment; filename=reports_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv')
    return response

@app.route('/export/pdf')
@permission_required('export_data')
def export_pdf():
    department = request.args.get('department', '')
    query = Report.query
    if department:
        query = query.filter(Report.department.ilike(f'%{department}%'))
    reports = query.order_by(Report.date_submitted.desc()).all()

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(200, 10, txt="Reports Export", align='C')
    pdf.ln(10)
    
    for r in reports:
        pdf.set_font("helvetica", style='B', size=10)
        pdf.cell(0, 10, txt=f"ID: {r.id} | {r.employee_name} ({r.department}) | Status: {r.status}", align='L')
        pdf.ln(8)
        pdf.set_font("helvetica", size=9)
        pdf.multi_cell(0, 6, txt=f"Submitted: {r.date_submitted.strftime('%Y-%m-%d %H:%M:%S')}\n{r.report_text}")
        pdf.ln(5)
        
    pdf_bytes = pdf.output()
    response = Response(bytes(pdf_bytes), mimetype='application/pdf')
    response.headers.set('Content-Disposition', f'attachment; filename=reports_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf')
    return response

@app.route('/export/docx')
@permission_required('export_data')
def export_docx():
    department = request.args.get('department', '')
    query = Report.query
    if department:
        query = query.filter(Report.department.ilike(f'%{department}%'))
    reports = query.order_by(Report.date_submitted.desc()).all()

    doc = Document()
    doc.add_heading('Reports Export', 0)
    
    for r in reports:
        p = doc.add_paragraph()
        p.add_run(f"ID: {r.id} | {r.employee_name} ({r.department}) | Status: {r.status}").bold = True
        doc.add_paragraph(f"Submitted: {r.date_submitted.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(r.report_text)
        doc.add_paragraph("-" * 40)
        
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    response = Response(f.read(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers.set('Content-Disposition', f'attachment; filename=reports_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    return response

@app.route('/audit-log')
@permission_required('manage_users') 
def audit_log():
    page = request.args.get('page', 1, type=int)
    pagination = db.session.query(Report, User).join(User, Report.user_id == User.id).order_by(Report.date_submitted.desc()).paginate(page=page, per_page=20, error_out=False)
    return render_template('audit_log.html', pagination=pagination)

@app.route('/admin/departments', methods=['GET', 'POST'])
@permission_required('manage_departments')
def admin_departments():
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        if name:
            if Department.query.filter_by(name=name).first():
                flash('Department already exists.', 'danger')
            else:
                dept = Department(name=name, description=description)
                db.session.add(dept)
                db.session.commit()
                flash('Department created.', 'success')
        return redirect(url_for('admin_departments'))
    departments = Department.query.all()
    return render_template('admin_departments.html', departments=departments)

@app.route('/admin/departments/<int:id>/edit', methods=['POST'])
@permission_required('manage_departments')
def edit_department(id):
    dept = db.session.get(Department, id)
    if dept:
        dept.name = request.form.get('name', dept.name)
        dept.description = request.form.get('description', dept.description)
        db.session.commit()
        flash('Department updated.', 'success')
    return redirect(url_for('admin_departments'))

@app.route('/admin/departments/<int:id>/delete', methods=['POST'])
@permission_required('manage_departments')
def delete_department(id):
    dept = db.session.get(Department, id)
    if dept:
        db.session.delete(dept)
        db.session.commit()
        flash('Department deleted.', 'success')
    return redirect(url_for('admin_departments'))

@app.route('/admin/users', methods=['GET', 'POST'])
@permission_required('manage_users')
def admin_users():
    if request.method == 'POST':
        action = request.form.get('action', 'update_role')
        if action == 'update_role':
            user_id = request.form.get('user_id')
            new_role = request.form.get('role')
            user = db.session.get(User, int(user_id))
            if user and user.id != current_user.id:
                user.role = new_role
                db.session.commit()
                flash(f"User {user.username} role updated to {new_role}.", 'success')
            elif user and user.id == current_user.id:
                flash("Cannot change your own role.", 'danger')
        return redirect(url_for('admin_users'))
    users = User.query.all()
    return render_template('admin_users.html', users=users, roles=ROLES.keys())

@app.route('/admin/users/create', methods=['POST'])
@permission_required('manage_users')
def admin_create_user():
    username = request.form.get('username')
    email = request.form.get('email')
    full_name = request.form.get('full_name')
    password = request.form.get('password')
    role = request.form.get('role', 'employee')
    if not username or not email or not password:
        flash('Username, email and password are required.', 'danger')
        return redirect(url_for('admin_users'))
    if User.query.filter_by(email=email).first():
        flash('Email already exists.', 'danger')
        return redirect(url_for('admin_users'))
    if User.query.filter_by(username=username).first():
        flash('Username already exists.', 'danger')
        return redirect(url_for('admin_users'))
    new_user = User(
        username=username, email=email, full_name=full_name,
        password=generate_password_hash(password), role=role
    )
    db.session.add(new_user)
    db.session.commit()
    flash(f'User {full_name or username} created successfully.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/<int:id>/edit', methods=['POST'])
@permission_required('manage_users')
def edit_user(id):
    user = db.session.get(User, id)
    if not user:
        flash('User not found.', 'danger')
        return redirect(url_for('admin_users'))
    user.full_name = request.form.get('full_name', user.full_name)
    user.email = request.form.get('email', user.email)
    user.username = request.form.get('username', user.username)
    new_password = request.form.get('new_password')
    if new_password:
        user.password = generate_password_hash(new_password)
    db.session.commit()
    flash(f'User {user.username} updated.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/users/<int:id>/delete', methods=['POST'])
@permission_required('manage_users')
def delete_user(id):
    if id == current_user.id:
        flash("You cannot delete yourself.", 'danger')
        return redirect(url_for('admin_users'))
    user = db.session.get(User, id)
    if user:
        db.session.delete(user)
        db.session.commit()
        flash("User deleted.", 'success')
    return redirect(url_for('admin_users'))

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    message = None
    if request.method == 'POST':
        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()
        message = 'If this email exists, please contact your system administrator to reset your password.'
    return render_template('forgot_password.html', message=message)

# All available permissions in the system
ALL_PERMISSIONS = [
    ('view_all_reports',  'View All Reports'),
    ('view_dept_reports', 'View Department Reports'),
    ('view_own_reports',  'View Own Reports'),
    ('submit_report',     'Submit Reports'),
    ('view_analytics',    'View Analytics'),
    ('approve_reports',   'Approve / Reject Reports'),
    ('delete_reports',    'Delete Reports'),
    ('export_data',       'Export Data (CSV/PDF/DOCX)'),
    ('manage_users',      'Manage Users & Roles'),
    ('manage_departments','Manage Departments'),
    ('manage_performance','Manage Performance Ledger'),
    ('manage_holidays',   'Manage Holidays'),
    ('manage_kpis',       'Manage KPIs & Entries'),
    ('manage_leaves',     'Manage Leave Requests'),
]

@app.route('/admin/roles', methods=['GET', 'POST'])
@permission_required('manage_users')
def admin_roles():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        selected_perms = request.form.getlist('permissions')
        if not name:
            flash('Role name is required.', 'danger')
            return redirect(url_for('admin_roles'))
        if name in ROLES:
            flash(f'"{name}" is a built-in role and cannot be overwritten.', 'danger')
            return redirect(url_for('admin_roles'))
        if CustomRole.query.filter_by(name=name).first():
            flash('A role with this name already exists.', 'danger')
            return redirect(url_for('admin_roles'))
        role = CustomRole(name=name, description=description, permissions=json.dumps(selected_perms))
        db.session.add(role)
        db.session.commit()
        flash(f'Role "{name}" created successfully.', 'success')
        return redirect(url_for('admin_roles'))
    custom_roles = CustomRole.query.all()
    return render_template('admin_roles.html', custom_roles=custom_roles,
                           all_permissions=ALL_PERMISSIONS, builtin_roles=ROLES)

@app.route('/admin/roles/<int:id>/edit', methods=['POST'])
@permission_required('manage_users')
def edit_role(id):
    role = db.session.get(CustomRole, id)
    if not role:
        flash('Role not found.', 'danger')
        return redirect(url_for('admin_roles'))
    role.description = request.form.get('description', role.description)
    selected_perms = request.form.getlist('permissions')
    role.permissions = json.dumps(selected_perms)
    db.session.commit()
    flash(f'Role "{role.name}" updated.', 'success')
    return redirect(url_for('admin_roles'))

@app.route('/admin/roles/<int:id>/delete', methods=['POST'])
@permission_required('manage_users')
def delete_role(id):
    role = db.session.get(CustomRole, id)
    if role:
        # Reset any users with this role to 'employee'
        User.query.filter_by(role=role.name).update({'role': 'employee'})
        db.session.delete(role)
        db.session.commit()
        flash(f'Role "{role.name}" deleted. Affected users reset to Employee.', 'success')
    return redirect(url_for('admin_roles'))

@app.route('/admin/holidays', methods=['GET', 'POST'])
@permission_required('manage_holidays')
def admin_holidays():
    if request.method == 'POST':
        date_str = request.form.get('date')
        description = request.form.get('description')
        if date_str:
            try:
                date_val = datetime.strptime(date_str, '%Y-%m-%d').date()
                if Holiday.query.filter_by(date=date_val).first():
                    flash('Holiday already exists for this date.', 'danger')
                else:
                    holiday = Holiday(date=date_val, description=description)
                    db.session.add(holiday)
                    db.session.commit()
                    flash('Holiday added.', 'success')
            except ValueError:
                flash('Invalid date format.', 'danger')
        return redirect(url_for('admin_holidays'))
    holidays = Holiday.query.order_by(Holiday.date.desc()).all()
    return render_template('admin_holidays.html', holidays=holidays)

@app.route('/admin/holidays/<int:id>/delete', methods=['POST'])
@permission_required('manage_holidays')
def delete_holiday(id):
    holiday = db.session.get(Holiday, id)
    if holiday:
        db.session.delete(holiday)
        db.session.commit()
        flash('Holiday removed.', 'success')
    return redirect(url_for('admin_holidays'))

@app.route('/admin/performance', methods=['GET', 'POST'])
@permission_required('manage_performance')
def admin_performance():
    if request.method == 'POST':
        employee_id = request.form.get('employee_id')
        entry_type = request.form.get('entry_type')
        description = request.form.get('description')
        if employee_id and description and entry_type:
            note = PerformanceNote(employee_id=employee_id, admin_id=current_user.id, entry_type=entry_type, description=description)
            db.session.add(note)
            db.session.commit()
            flash('Performance entry logged.', 'success')
        return redirect(url_for('admin_performance'))
        
    employee_filter = request.args.get('employee_id', '')
    query = PerformanceNote.query
    if employee_filter:
        query = query.filter_by(employee_id=employee_filter)
        
    notes = query.order_by(PerformanceNote.date_logged.desc()).all()
    employees = User.query.filter_by(role='employee').all()
    return render_template('admin_performance.html', notes=notes, employees=employees, employee_filter=employee_filter)

@app.route('/admin/performance/<int:id>/delete', methods=['POST'])
@permission_required('manage_performance')
def delete_performance(id):
    note = db.session.get(PerformanceNote, id)
    if note:
        db.session.delete(note)
        db.session.commit()
        flash('Performance log deleted.', 'success')
    return redirect(url_for('admin_performance'))

@app.route('/admin/kpis', methods=['GET', 'POST'])
@permission_required('manage_kpis')
def admin_kpis():
    if request.method == 'POST':
        department_id = request.form.get('department_id')
        metric_name = request.form.get('metric_name')
        metric_type = request.form.get('metric_type')
        target_value = request.form.get('target_value')
        description = request.form.get('description')
        
        if department_id and metric_name:
            target = float(target_value) if target_value else None
            kpi = KPIDefinition(
                department_id=department_id, 
                metric_name=metric_name, 
                metric_type=metric_type, 
                target_value=target,
                description=description
            )
            db.session.add(kpi)
            db.session.commit()
            flash('KPI Definition added.', 'success')
        return redirect(url_for('admin_kpis'))
        
    kpis = KPIDefinition.query.all()
    departments = Department.query.all()
    return render_template('admin_kpis.html', kpis=kpis, departments=departments)

@app.route('/admin/kpis/<int:id>/delete', methods=['POST'])
@permission_required('manage_kpis')
def delete_kpi(id):
    kpi = db.session.get(KPIDefinition, id)
    if kpi:
        db.session.delete(kpi)
        db.session.commit()
        flash('KPI Definition deleted.', 'success')
    return redirect(url_for('admin_kpis'))

@app.route('/manager/kpi_log', methods=['GET', 'POST'])
@permission_required('manage_kpis')
def manager_kpi_log():
    if request.method == 'POST':
        employee_id = request.form.get('employee_id')
        date_str = request.form.get('date')
        
        if employee_id and date_str:
            date_val = datetime.strptime(date_str, '%Y-%m-%d').date()
            logged_count = 0
            for key, value in request.form.items():
                if key.startswith('kpi_') and value.strip() != '':
                    kpi_id = int(key.replace('kpi_', ''))
                    actual_val = float(value)
                    
                    # Check if already logged for this date
                    existing_log = DailyKPILog.query.filter_by(
                        employee_id=employee_id, 
                        date=date_val, 
                        kpi_definition_id=kpi_id
                    ).first()
                    
                    if existing_log:
                        existing_log.actual_value = actual_val
                    else:
                        new_log = DailyKPILog(
                            employee_id=employee_id,
                            manager_id=current_user.id,
                            date=date_val,
                            kpi_definition_id=kpi_id,
                            actual_value=actual_val
                        )
                        db.session.add(new_log)
                    logged_count += 1
                    
            if logged_count > 0:
                db.session.commit()
                flash(f'Successfully logged {logged_count} KPI metrics.', 'success')
            else:
                flash('No KPI metrics were filled out.', 'warning')
                
        return redirect(url_for('manager_kpi_log'))
        
    employees = User.query.filter_by(role='employee').all()
    departments = Department.query.all()
    kpis = KPIDefinition.query.all()
    recent_logs = DailyKPILog.query.order_by(DailyKPILog.date.desc()).limit(20).all()
    return render_template('manager_kpi_log.html', employees=employees, departments=departments, kpis=kpis, recent_logs=recent_logs)

@app.route('/leaves', methods=['GET', 'POST'])
@login_required
def leave_requests():
    if request.method == 'POST':
        start_date_str = request.form.get('start_date')
        end_date_str = request.form.get('end_date')
        reason = request.form.get('reason')
        if start_date_str and end_date_str and reason:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
            if end_date < start_date:
                flash('End date must be after start date.', 'danger')
            else:
                leave = LeaveRequest(user_id=current_user.id, start_date=start_date, end_date=end_date, reason=reason)
                db.session.add(leave)
                db.session.commit()
                flash('Leave request submitted successfully.', 'success')
        return redirect(url_for('leave_requests'))
        
    if current_user.has_permission('manage_leaves') and current_user.role != 'employee':
        leaves = LeaveRequest.query.order_by(LeaveRequest.created_at.desc()).all()
    else:
        leaves = LeaveRequest.query.filter_by(user_id=current_user.id).order_by(LeaveRequest.created_at.desc()).all()
        
    return render_template('leave_requests.html', leaves=leaves)

@app.route('/leaves/<int:id>/status', methods=['POST'])
@permission_required('manage_leaves')
def update_leave_status(id):
    if current_user.role == 'employee':
        flash('Permission denied.', 'danger')
        return redirect(url_for('leave_requests'))
        
    leave = db.session.get(LeaveRequest, id)
    if leave:
        status = request.form.get('status')
        if status in ['Approved', 'Rejected']:
            leave.status = status
            db.session.commit()
            flash(f'Leave request {status.lower()}.', 'success')
    return redirect(url_for('leave_requests'))
    
@app.route('/admin/migrate-db')
@login_required
def migrate_db():
    if current_user.role != 'admin':
        return "Admin only", 403
    try:
        from sqlalchemy import text
        # Attempt to add the column. Try/Except in case it exists.
        db.session.execute(text("ALTER TABLE kpi_definition ADD COLUMN IF NOT EXISTS description TEXT"))
        # IF NOT EXISTS is Postgres specific. For SQLite we handle the error.
        db.session.commit()
        return "Migration successful (Postgres style)!"
    except Exception as e:
        db.session.rollback()
        try:
            db.session.execute(text("ALTER TABLE kpi_definition ADD COLUMN description TEXT"))
            db.session.commit()
            return "Migration successful (General style)!"
        except Exception as e2:
            return f"Migration failed or already done. Error: {e2}"

@app.route('/admin/seed-kpis')
@login_required
def seed_kpis_route():
    if current_user.role != 'admin':
        return "Admin only", 403
    try:
        # Create Departments
        depts_data = [
            "Sales Team", "Clearance Team", "After-Use Team", 
            "Customer Issue Management", "Call Confirmation", 
            "Order Invoice Creation Team", "Warehouse & Fulfillment", 
            "Delivery Team", "Operation Manager"
        ]
        
        dept_map = {}
        for name in depts_data:
            dept = Department.query.filter_by(name=name).first()
            if not dept:
                dept = Department(name=name)
                db.session.add(dept)
                db.session.commit()
            dept_map[name] = dept.id

        # KPI Data
        kpis = [
            ("Sales Team", "Response Professionalism", "percentage", 100, "Review chat history for tone and accuracy."),
            ("Sales Team", "Message Success Rate", "percentage", 100, "Weekly count of correct replies. Target: 0% error."),
            ("Sales Team", "Follow-up Conversion", "percentage", 25, "Target: 20%-30% conversion of pending customers."),
            ("Sales Team", "Cross-sell Ratio", "percentage", 20, "At least 15-20% orders should have related products added."),
            ("Sales Team", "Data Error Rate", "percentage", 0, "Parcels returned due to data entry errors."),
            ("Sales Team", "Note Compliance", "percentage", 100, "Percentage of orders with text notes (Address + List)."),
            ("Sales Team", "Stock-out Confirmation", "percentage", 0, "Orders confirmed for out-of-stock items."),
            ("Sales Team", "Daily Reporting", "percentage", 100, "On-time and accurate EOD report submission."),
            ("Clearance Team", "Message Clearance Speed", "number", 75, "Target: 60-90 messages per hour."),
            ("Clearance Team", "Priority Handling", "percentage", 100, "Handling labels like 'Ordered Customer' first."),
            ("Clearance Team", "Information Accuracy", "percentage", 100, "Zero pricing or info errors from Doc."),
            ("Clearance Team", "Push/Motivation Rate", "percentage", 50, "Percentage of price queries pushed towards sales."),
            ("Clearance Team", "Meta Response Rate", "percentage", 90, "Overall meta response rate target."),
            ("After-Use Team", "First Response Time", "number", 5, "Target: Within 5 minutes (max 10)."),
            ("After-Use Team", "Daily Case Volume", "number", 35, "Target: 25-40+ cases handled per day."),
            ("After-Use Team", "Documentation Rate", "percentage", 100, "Notes for Problem + Solution + Follow-up."),
            ("After-Use Team", "Reopen Case Rate", "percentage", 15, "Target: Below 15% (same issue returning)."),
            ("After-Use Team", "Sensitive Follow-up", "percentage", 100, "100% sensitive cases followed up within 1 week."),
            ("Customer Issue Management", "Resolution Time", "number", 24, "Target: 100% issues updated/resolved within 24h."),
            ("Customer Issue Management", "Exchange Processing", "number", 0, "Delay in notifying delivery team for exchange."),
            ("Customer Issue Management", "Comment Response Rate", "percentage", 100, "100% coverage of group/page comments."),
            ("Customer Issue Management", "Backlog Prevention", "percentage", 100, "Zero pending receipts/invoices at end of day."),
            ("Call Confirmation", "SOP Compliance", "percentage", 100, "Script, behavior, and verification accuracy."),
            ("Call Confirmation", "Data Accuracy", "percentage", 100, "Zero errors in product, size, or price during call."),
            ("Call Confirmation", "Advance Compliance", "percentage", 100, "Collecting advance for 10k+ or high-risk orders."),
            ("Call Confirmation", "Tagging Accuracy", "percentage", 100, "Correct usage of Confirmed, Mirpur, Hold, etc."),
            ("Order Invoice Creation Team", "Profile Accuracy", "percentage", 100, "Zero duplicate profiles or data entry errors."),
            ("Order Invoice Creation Team", "Stock Alert Speed", "number", 0, "Immediate reporting of stock-outs to sales."),
            ("Order Invoice Creation Team", "Daily Receipt Count", "number", 150, "Minimum 150 receipts cut per day."),
            ("Order Invoice Creation Team", "Advance Policy Adherence", "percentage", 100, "Ensuring advance for return-history orders."),
            ("Warehouse & Fulfillment", "Picking Accuracy", "percentage", 100, "Zero errors in product, shade, or quantity."),
            ("Warehouse & Fulfillment", "Packaging Safety", "percentage", 100, "100% bubble wrap and secure sealing."),
            ("Warehouse & Fulfillment", "Dispatch Safety", "percentage", 100, "Zero 'Hold' or 'Cancel' parcels shipped."),
            ("Warehouse & Fulfillment", "Fulfillment Volume", "number", 1000, "Target: 1000 parcels per day."),
            ("Warehouse & Fulfillment", "Resolution TAT", "number", 48, "Average 24-48h to resolve stock/hold issues."),
            ("Delivery Team", "Delivery Success Rate", "percentage", 95, "Target 95%-100% successful delivery."),
            ("Delivery Team", "Pickup Accuracy", "percentage", 100, "Zero errors or damage during supplier pickup."),
            ("Delivery Team", "Cash Handling", "percentage", 100, "Zero mismatch in cash collection reports."),
            ("Operation Manager", "SOP Compliance Audit", "percentage", 98, "Random audit scores across all teams."),
            ("Operation Manager", "Response Time Improvement", "percentage", 10, "MoM improvement in response times."),
            ("Operation Manager", "Reporting Timeliness", "percentage", 100, "On-time submission of daily/weekly reports.")
        ]
        
        for dept_name, metric, mtype, target, desc in kpis:
            dept_id = dept_map[dept_name]
            existing = KPIDefinition.query.filter_by(department_id=dept_id, metric_name=metric).first()
            if not existing:
                kpi = KPIDefinition(
                    department_id=dept_id,
                    metric_name=metric,
                    metric_type=mtype,
                    target_value=target,
                    description=desc
                )
                db.session.add(kpi)
        
        db.session.commit()
        return "Seeding successful!"
    except Exception as e:
        return f"Seeding failed: {e}"

@app.route('/api/reports')
def api_reports():
    page = request.args.get('page', 1, type=int)
    pagination = Report.query.order_by(Report.date_submitted.desc()).paginate(page=page, per_page=10, error_out=False)
    data = []
    for r in pagination.items:
        data.append({
            'id': r.id,
            'employee_name': r.employee_name,
            'department': r.department,
            'report_text': r.report_text,
            'status': r.status,
            'date_submitted': r.date_submitted.isoformat()
        })
    return jsonify({
        'reports': data,
        'total': pagination.total,
        'pages': pagination.pages,
        'current_page': pagination.page
    })

@app.route('/api/analytics/employee/<employee_name>')
@permission_required('view_analytics')
def api_employee_analytics(employee_name):
    user = User.query.filter((User.full_name == employee_name) | (User.username == employee_name)).first()
    reports = Report.query.filter_by(employee_name=employee_name).all()
    
    total = len(reports)
    approved = sum(1 for r in reports if r.status == 'Approved')
    rejected = sum(1 for r in reports if r.status == 'Rejected')
    submitted = sum(1 for r in reports if r.status == 'Submitted')
    
    seven_days_ago = datetime.now(timezone.utc) - timedelta(days=7)
    recent_reports = [r for r in reports if r.date_submitted >= seven_days_ago]
    
    trend_dict = {}
    for i in range(7, -1, -1):
        d = (datetime.now(timezone.utc) - timedelta(days=i)).date().strftime('%Y-%m-%d')
        trend_dict[d] = 0
        
    for r in recent_reports:
        d_str = r.date_submitted.date().strftime('%Y-%m-%d')
        if d_str in trend_dict:
            trend_dict[d_str] += 1
            
    kpi_data = []
    if user:
        thirty_days_ago = datetime.now(timezone.utc).date() - timedelta(days=30)
        logs = DailyKPILog.query.filter(DailyKPILog.employee_id == user.id, DailyKPILog.date >= thirty_days_ago).order_by(DailyKPILog.date.asc()).all()
        kpi_dict = {}
        for log in logs:
            k_id = log.kpi_definition_id
            if k_id not in kpi_dict:
                kpi_dict[k_id] = {
                    'name': log.kpi_definition.metric_name,
                    'target': log.kpi_definition.target_value,
                    'actuals': [],
                    'dates': []
                }
            kpi_dict[k_id]['actuals'].append(log.actual_value)
            kpi_dict[k_id]['dates'].append(log.date.strftime('%Y-%m-%d'))
        kpi_data = list(kpi_dict.values())
            
    return jsonify({
        'employee_name': employee_name,
        'total_reports': total,
        'status_breakdown': {
            'Approved': approved,
            'Rejected': rejected,
            'Submitted': submitted
        },
        'trend_labels': list(trend_dict.keys()),
        'trend_data': list(trend_dict.values()),
        'kpi_data': kpi_data
    })

with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(debug=os.environ.get('FLASK_DEBUG', 'True') == 'True', port=int(os.environ.get('PORT', 5000)))
